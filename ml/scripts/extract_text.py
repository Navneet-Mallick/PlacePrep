"""
Extract plain text from PDF or DOCX resume files.

Handles:
- Multi-page PDFs
- Encrypted PDFs (returns empty with error message)
- DOCX with tables and formatting
- Null byte cleanup
- Encoding issues

Usage:
    python ml/scripts/extract_text.py path/to/resume.pdf
"""

import sys
from pathlib import Path

from docx import Document
from PyPDF2 import PdfReader


def extract_from_pdf(path: Path) -> str:
    """Extract text from PDF with robust error handling"""
    try:
        reader = PdfReader(str(path))
        
        # Check for encryption
        if reader.is_encrypted:
            try:
                reader.decrypt("")  # Try empty password
            except Exception:
                print("Warning: PDF is encrypted and cannot be read", file=sys.stderr)
                return ""
        
        if not reader.pages:
            return ""
        
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    # Clean null bytes and control characters
                    text = text.replace('\x00', '')
                    text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                    if text.strip():
                        pages.append(text.strip())
            except Exception as e:
                print(f"Warning: Failed to extract page {i+1}: {e}", file=sys.stderr)
                continue
        
        result = "\n\n".join(pages)
        
        # Final cleanup
        # Remove excessive whitespace
        import re
        result = re.sub(r'\n{3,}', '\n\n', result)
        result = re.sub(r' {3,}', ' ', result)
        
        return result.strip()
        
    except Exception as e:
        print(f"Error reading PDF: {e}", file=sys.stderr)
        return ""


def extract_from_docx(path: Path) -> str:
    """Extract text from DOCX including tables"""
    try:
        document = Document(str(path))
        
        parts = []
        
        # Extract paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        
        # Extract tables (often contain skills, education)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        
        result = "\n".join(parts)
        
        # Clean null bytes
        result = result.replace('\x00', '')
        
        return result.strip()
        
    except Exception as e:
        print(f"Error reading DOCX: {e}", file=sys.stderr)
        return ""


def extract_text(path: Path) -> str:
    """Extract text from resume file (PDF or DOCX)"""
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return extract_from_pdf(path)
    elif suffix == ".docx":
        return extract_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use PDF or DOCX.")


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python ml/scripts/extract_text.py path/to/resume.pdf")
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"File not found: {path}")
        sys.exit(1)

    text = extract_text(path)
    if text:
        print(f"Extracted {len(text)} characters")
        print(text[:500])
    else:
        print("No text could be extracted from the file")


if __name__ == "__main__":
    main()
